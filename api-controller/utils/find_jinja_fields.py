import re, os
from docx import Document

def get_all_jinja_fields(doc_location):
	print('Fetching jinja fields from template:',doc_location)
	template_doc = Document(doc_location)
	jinja_field_regex = "{{[^{}]+}}"

	set_of_jinja_fields = set()
	# search in paragraphs
	for para in template_doc.paragraphs:
		list_of_matches = re.findall(jinja_field_regex, para.text)

		for jinja_field in list_of_matches:
			set_of_jinja_fields.add(jinja_field[2:][:-2])

	# search in tables
	for table in template_doc.tables:
		for row in table.rows:
			for cell in row.cells:
				list_of_matches = re.findall(jinja_field_regex, cell.text)
				for jinja_field in list_of_matches:
					set_of_jinja_fields.add(jinja_field[2:][:-2])

	#print(list(set_of_jinja_fields))
	return list(set_of_jinja_fields)

#get_all_jinja_fields(os.path.abspath("../template/user_a/sample.docx"))