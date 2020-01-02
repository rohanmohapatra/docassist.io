from flask import Blueprint,Response, jsonify, request
#from dataaccess.get_functions import *
from flask_cors import CORS, cross_origin
import json
import time
import datetime
import os
from flask import current_app as app
from dataccess.data_getfunctions import get_client_email_by_doc_name
from dataccess.templates_setfunctions import add_template_to_db, set_jinja_fields
from dataccess.templates_getfunctions import get_templates_for_user, get_template_by_id
from flask_cors import CORS, cross_origin
from utils.aes import encrypt,decrypt

from utils.utils import allowed_file_extensions
from werkzeug.utils import secure_filename
import yagmail

import spacy
import docx 
import difflib
import mammoth
from html2docx import html2docx
nlp = spacy.load("en_core_web_sm")


autotemplate_view = Blueprint('autotemplate_view',__name__)


def nlp_engine(texts):
    ent_texts_labels = []
    for doc in nlp.pipe(texts): #Processing texts as a stream and buffering in batches, instead of one-by-one
        entities = doc.ents
        for ent in entities:
            if(ent.text!='' and ent.text!='___________________________________________________'):
                ent_texts_labels.append((ent.text, ent.label_))
    return ent_texts_labels
def replace_string(filename, dict_fields):
    doc = docx.Document(filename)
    print(dict_fields)
    for p in doc.paragraphs:
        para_text=p.text
        for tag,value in dict_fields.items():
        
            if tag in para_text and para_text!='':
                inline = p.runs
            # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    print("Tag : ",tag, "\nInline text : ",inline[i].text)
                    if tag in inline[i].text:
                        text = inline[i].text.replace(tag, '{{'+value+'}}')
                        inline[i].text = text
                    if inline[i].text in tag:
                        text = inline[i].text.replace(inline[i].text,'{{'+value+'}}')
                        inline[i].text = text
                print("Updated text : ",p.text,"\n\n")
        #print(tag,"and",para_text)

    doc.save(os.path.join(app.config['AUTOTEMPLATE_UPLOAD_FOLDER'],'template.docx'))
    return 1

@autotemplate_view.route("/upload/", methods=['POST'])
@cross_origin()
def upload_docx():
    ALLOWED_EXTENSIONS = set(['docx'])
    if request.method == 'POST':
        #print(request.files.to_dict(flat=False))
        
        #Check if my inout has those fields
        if 'autotemplate' not in request.files:
            return Response(status=400)
        
        template = request.files['autotemplate']

        if template.filename == '':
            return Response(status=400)

        #If allowed extensions save the Docx
        if template and allowed_file_extensions(template.filename, ALLOWED_EXTENSIONS):
            template_filename = secure_filename(template.filename)
            try:
                #template_id = add_template_to_db(template_filename,str(os.path.join(app.config['UPLOAD_FOLDER'], template_filename)))
                template.save(os.path.join(app.config['AUTOTEMPLATE_UPLOAD_FOLDER'], template_filename))
            except Exception as e:
                print("Error:",e)
                print("Something went wrong")
                return Response(status=409)   

        else:
            return Response('["file types not supported"]', status=400)

        #Make a call to the string processing module
        data = dict()
        data['document_name'] = template_filename
        return jsonify(data)
    else:
        Response(status=405)

@autotemplate_view.route("/generate/", methods=["POST"])
@cross_origin()
def generate():
    json_data = request.get_json(force=True)
    document_1 = json_data["document_1"]
    document_2 = json_data["document_2"]
    doc1 = docx.Document(os.path.join(app.config['AUTOTEMPLATE_UPLOAD_FOLDER'],document_1))
    doc2 = docx.Document(os.path.join(app.config['AUTOTEMPLATE_UPLOAD_FOLDER'],document_2))
    #Get the difference
    paras_doc1 = []
    paras_doc2 = []

    for i in doc1.paragraphs:
        if(i.text!=' ' and i.text!=''):
            paras_doc1.append(i.text)

    for i in doc2.paragraphs:
        if(i.text!=' ' and i.text!='' ):
            paras_doc2.append(i.text)

    
    entities_doc1 = nlp_engine(paras_doc1)
    entities_doc2 = nlp_engine(paras_doc2)

    diff_tool = difflib.Differ()
    diff_list=[]
    d = difflib.Differ()
    for i in range(len(paras_doc1)):
        t1 = paras_doc1[i]
        t2 = paras_doc2[i]
        res = list(d.compare(t1,t2))
        diff_list.append(res)
    result = list(d.compare(paras_doc1, paras_doc2))

    doc1_diffs=[]
    doc2_diffs=[]
    for i in result:
        if(i[0] =='-'):
            doc1_diffs.append(i[1:])
        elif(i[0] == '+'):
            doc2_diffs.append(i[1:])

    entities_doc1 = nlp_engine(doc1_diffs)
    entities_doc2 = nlp_engine(doc2_diffs)

    dic={}
    for i in entities_doc1:
        dic.update({i[0]:i[1]})
    print(dic)

    replace_string(os.path.join(app.config['AUTOTEMPLATE_UPLOAD_FOLDER'],document_1),dic)

    data = dict()
    data['generated_document_name'] ="template.docx"
    return jsonify(data)


@autotemplate_view.route("/edit/<generated_document_name>/", methods=["GET"])
@cross_origin()
def edit_generated(generated_document_name):
    document = os.path.join(app.config['AUTOTEMPLATE_UPLOAD_FOLDER'], generated_document_name)
    with open(document, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value # The generated HTML
        #print(html)
        return html

@autotemplate_view.route("/save/", methods=['POST'])
@cross_origin()
def save_generated_template():
    json_data = request.get_json(force=True)
    #print(json_data)
    html = json_data["html"]
    template_filename = json_data["template_filename"]
    buf = html2docx(html, title="Auto Generated Template")
    #print(location)
    location = os.path.join(app.config['UPLOAD_FOLDER'],template_filename)
    with open(location, "wb") as fp:
        fp.write(buf.getvalue())
        template_id = add_template_to_db(template_filename,str(os.path.join(app.config['UPLOAD_FOLDER'], template_filename)))
        #template.save(os.path.join(app.config['UPLOAD_FOLDER'], template_filename))
        # Now add the jinja-fields in the template to the template document in MongoDB
        template_location = os.path.join(app.config['UPLOAD_FOLDER'], template_filename)
        modified_count = set_jinja_fields(template_location, template_id)

        if(modified_count):
            print("Added jinja-fields for",str(modified_count),"template")
            return Response(status=200)
        else:
            print("Failed to add jinja fields to mongo.db.template")
            return Response(status=409)
    return Response(status=200)