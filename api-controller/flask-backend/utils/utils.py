import re, os
from docx import Document

def get_all_jinja_fields(doc_location):
	print('Fetching jinja fields from template:',doc_location)
	template_doc = Document(doc_location)
	jinja_field_regex = "{{[^{}]+}}"

	
	set_of_jinja_fields = set()
	set_of_sub_templates = set()

	# search in paragraphs
	for para in template_doc.paragraphs:
		list_of_matches = re.findall(jinja_field_regex, para.text)

		for jinja_field in list_of_matches:
			if (jinja_field[2:4]=="p "):
				set_of_sub_templates.add(jinja_field[4:][:-2]+'.docx')
			else:
				set_of_jinja_fields.add(jinja_field[2:][:-2])

	# search in tables
	for table in template_doc.tables:
		for row in table.rows:
			for cell in row.cells:
				list_of_matches = re.findall(jinja_field_regex, cell.text)
				for jinja_field in list_of_matches:
					if (jinja_field[2:4]=="p "):
						set_of_sub_templates.add(jinja_field[4:][:-2]+'.docx')
					else:
						set_of_jinja_fields.add(jinja_field[2:][:-2])

	return list(set_of_jinja_fields), list(set_of_sub_templates)

def allowed_file_extensions(filename, ALLOWED_EXTENSIONS):
	return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def reshape(src_list, shape):
	'''
	shape is a tuple
	assume (-1,1)
	'''
	i=0
	new_list=[]
	while i<len(src_list):
		new_list.append(src_list[i:i+shape[1]])
		i+=shape[1]
	return new_list